#!/usr/bin/env python
import os
import shutil
import smtplib
from cgi import parse_header
from datetime import date
from email.message import EmailMessage
from functools import wraps
from io import BytesIO

import requests
import logging

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from flask import Flask, Response, abort, request


app = Flask(__name__)
logging.basicConfig(filename='error.log', level=logging.ERROR)

secret_key = os.environ.get("SECRET_KEY", "foo")
api_key = os.environ.get("API_KEY")

from_email = os.environ.get("SENDER_ADDRESS")
to_email = os.environ.get("RECIPIENT_ADDRESS")
copy_emails = os.environ.get("COPY_ADDRESSES", ",").split(",")

host = os.environ.get("SMTP_HOST")
port = os.environ.get("SMTP_PORT")
username = os.environ.get("SMTP_USERNAME")
password = os.environ.get("SMTP_PASSWORD")


def require_apikey(view_function):
    """Function to confirm request contains a valid API key"""

    @wraps(view_function)
    def decorated_function(*args, **kwargs):
        if request.headers.get("Authorization", "") == f"Token {api_key}":
            return view_function(*args, **kwargs)
        abort(401)

    return decorated_function


@app.route("/", methods=["GET"])
def index():
    return "alive"


@app.route("/submit", methods=["GET", "POST"])
@require_apikey
def submit():
    fields = ["name", "email", "supplier", "date", "amount", "description", "receipt_id"]
    if request.method == "GET":
        s = '<form method="post">'
        s = s + "".join([f'{f}: <input name="{f}" /> <br />' for f in fields])
        s = s + "<input type='submit' /></form>"
        return s
    if request.method == "POST":
        if not all([f in request.form for f in fields]):
            return Response(status=400)
        doc_stream, receipt_stream = generate_report(request.form)
        send_report(doc_stream, receipt_stream, request.form, request.args.get("is_test", False))
        return "message sent!"


def send_report(doc_stream, receipt_stream, data, is_test=False):
    msg = EmailMessage()

    msg[
        "Subject"
    ] = f'Penn Labs Expense report for purchase from {data["supplier"]} on {data["date"]}'

    msg["From"] = from_email
    msg["To"] = to_email if not is_test else from_email
    msg["Cc"] = copy_emails

    msg.set_content(
        "Hi,\n\nWe're attaching a completed Penn Labs expense report.\n\nBest,\nThe Penn Labs Directors"  # noqa
    )

    msg.add_attachment(
        doc_stream.read(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{data['supplier']} Expense Report (Receipt Attached).docx",
    )

    r, t = receipt_stream
    if r is not None:
        maint, subt = t.split("/")
        msg.add_attachment(r.read(), maintype=maint, subtype=subt, filename=f"Receipt.{subt}")

    with smtplib.SMTP(host, 587) as server:
        server.login(username, password)
        server.send_message(msg)
        return True


def generate_report(data):
    document = Document("REPORT_TEMPLATE.docx")
    tables = document.tables

    tables[0].cell(1, 0).text = data["name"]
    tables[0].cell(1, 1).text = data["email"]
    tables[0].cell(1, 2).text = date.today().strftime("%-m/%-d/%Y")

    tables[2].cell(1, 0).text = data["supplier"]
    tables[2].cell(1, 1).text = data["date"]
    tables[2].cell(1, 2).text = data["amount"]

    tables[3].cell(1, 0).text = data["description"]

    receipt_id = data["receipt_id"].split("id=")[1]
    url = f"https://drive.google.com/uc?id={receipt_id}&export=download"

    d = BytesIO()
    i = BytesIO()

    paragraph = document.add_paragraph()
    run = paragraph.add_run("Receipt:")
    run.bold = True
    run.font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    receipt_stream = None
    try:
        with requests.get(url, stream=True) as r:
            mtype, _ = parse_header(r.headers.get("content-type"))
            img = BytesIO(r.content)
            document.add_picture(img, width=Pt(400))

            shutil.copyfileobj(r.raw, i)
            i.seek(0)
            receipt_stream = (i, mtype)
    except Exception as e:
        logging.error(f"Error while adding receipt image: {str(e)}")
        receipt_stream = (None, None)

    document.save(d)
    d.seek(0)

    return d, receipt_stream

if __name__ == "__main__":
    app.run()
