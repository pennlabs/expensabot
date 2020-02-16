#!/user/bin/python
from datetime import date
import requests
import shutil
import re

from PIL import Image

from docx import Document
from docx.shared import Inches


def generate_receipt(data):
    document = Document("expense/REPORT_TEMPLATE.docx")
    tables = document.tables

    tables[0].cell(1, 0).text = data["name"]
    tables[0].cell(1, 1).text = data["email"]
    tables[0].cell(1, 2).text = date.today().strftime("%-m/%d/%Y")

    tables[2].cell(1, 0).text = data["supplier"]
    tables[2].cell(1, 1).text = data["date"]
    tables[2].cell(1, 2).text = data["amount"]

    tables[3].cell(1, 0).text = data["description"]

    local_filename = "expense/receipt"

    url = None
    id_match = re.search(r"id=([0-9a-zA-Z]+)", data["receipt"])
    print(id_match)
    if id_match is not None:
        url = f"https://drive.google.com/uc?id={id_match.group(1)}&export=download"

    if url is None:
        print("couldn't extract receipt. exiting...")
        exit(1)

    with requests.get(url, stream=True) as r:
        with open(local_filename, "wb") as f:
            shutil.copyfileobj(r.raw, f)
    Image.open(local_filename).save(f"{local_filename}.jpg")
    document.add_picture(f"{local_filename}.jpg", height=Inches(8))

    document.save("output.docx")


if __name__ == "__main__":
    generate_receipt(
        {
            "name": "Armaan",
            "email": "armaan@example.com",
            "supplier": "yo",
            "date": "2/2/2020",
            "amount": "$12.75",
            "description": "blah blah blah",
            "receipt": "https://drive.google.com/open?id=1i16kjXWk0zEnaiyiBHVSSSgoQn2GQw9Q",
        }
    )
